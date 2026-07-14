from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile, ZipFile

from docx import Document
from pypdf import PdfReader


class DocumentParseError(ValueError):
    pass


@dataclass(frozen=True, slots=True)
class ParsedSection:
    text: str
    page_number: int | None = None


SUPPORTED_EXTENSIONS = frozenset({".pdf", ".docx", ".txt", ".md"})


def _clean_text(value: str) -> str:
    value = value.replace("\x00", " ").replace("\r\n", "\n").replace("\r", "\n")
    value = re.sub(r"[\t\u00a0]+", " ", value)
    value = re.sub(r"[ ]{2,}", " ", value)
    value = re.sub(r"\n[ ]+", "\n", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def _parse_pdf(data: bytes) -> list[ParsedSection]:
    try:
        reader = PdfReader(BytesIO(data), strict=False)
        if reader.is_encrypted and reader.decrypt("") == 0:
            raise DocumentParseError("Password-protected PDF files are not supported")
        sections = []
        for index, page in enumerate(reader.pages, start=1):
            text = _clean_text(page.extract_text() or "")
            if text:
                sections.append(ParsedSection(text=text, page_number=index))
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError("The PDF could not be parsed") from exc
    if not sections:
        raise DocumentParseError(
            "The PDF contains no extractable text; scanned files require OCR before ingestion"
        )
    return sections


def _parse_docx(data: bytes) -> list[ParsedSection]:
    try:
        document = Document(BytesIO(data))
        blocks: list[str] = []
        blocks.extend(p.text for p in document.paragraphs if p.text.strip())
        for table in document.tables:
            for row in table.rows:
                cells = [_clean_text(cell.text) for cell in row.cells]
                line = " | ".join(cell for cell in cells if cell)
                if line:
                    blocks.append(line)
        text = _clean_text("\n\n".join(blocks))
    except Exception as exc:
        raise DocumentParseError("The DOCX file could not be parsed") from exc
    if not text:
        raise DocumentParseError("The DOCX file contains no extractable text")
    return [ParsedSection(text=text)]


def _parse_text(data: bytes) -> list[ParsedSection]:
    try:
        text = data.decode("utf-8-sig")
    except UnicodeDecodeError:
        try:
            text = data.decode("windows-1252")
        except UnicodeDecodeError as exc:
            raise DocumentParseError("Text files must use UTF-8 or Windows-1252 encoding") from exc
    text = _clean_text(text)
    if not text:
        raise DocumentParseError("The text file is empty")
    return [ParsedSection(text=text)]


def parse_document(data: bytes, filename: str) -> list[ParsedSection]:
    if not data:
        raise DocumentParseError("The uploaded file is empty")
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        allowed = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise DocumentParseError(f"Unsupported file type; accepted extensions are {allowed}")
    if extension == ".pdf":
        if not data.lstrip().startswith(b"%PDF-"):
            raise DocumentParseError("File content does not match the .pdf extension")
        return _parse_pdf(data)
    if extension == ".docx":
        try:
            with ZipFile(BytesIO(data)) as archive:
                names = set(archive.namelist())
                if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                    raise DocumentParseError("File content does not match the .docx extension")
        except BadZipFile as exc:
            raise DocumentParseError("File content does not match the .docx extension") from exc
        return _parse_docx(data)
    return _parse_text(data)


def sections_text(sections: list[ParsedSection]) -> str:
    return "\n\n".join(section.text for section in sections).strip()